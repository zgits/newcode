import requests
from bs4 import BeautifulSoup
import re
import docx

def get_html(url):
    headers = {
        'User-Agent': 'Mozilla/5.0(Macintosh; Intel Mac OS X 10_11_4)\
        AppleWebKit/537.36(KHTML, like Gecko) Chrome/52 .0.2743. 116 Safari/537.36'

    }  # 模拟浏览器访问
    response = requests.get(url, headers=headers)  # 请求访问网站
    html = response.text  # 获取网页源码
    return html  # 返回网页源码
file=docx.Document()
for i in range(1,120):
    soup = BeautifulSoup(get_html('https://www.nowcoder.com/ta/review-java/review?page='+str(i)), 'html.parser')  # 初始化BeautifulSoup库,并设置解析器
    #print(get_html('https://www.nowcoder.com/ta/review-java/review?page=2'))
    question=soup.find_all('div',attrs={'class':'final-question'})
    answer = soup.find_all('div',attrs={'class':'design-answer-box'})
    file.add_paragraph(str(i)+'.'+question[0].string)
    file.add_paragraph('答：'+str(answer[0]))
file.save("C:\\Users\\幻夜~星辰\\Desktop\work\\newcode.docx")








